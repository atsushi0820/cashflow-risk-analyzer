"""
資金繰りリスク可視化ツール - 完全版（Word/Excel出力対応）
Phase 3（長期運転資金）→ Phase 2（ストレステスト）→ レポート出力
"""

import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from datetime import datetime
from io import BytesIO

# Word出力用
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Excel出力用
try:
    import openpyxl
    from openpyxl.styles import Font, Alignment, PatternFill
    from openpyxl.chart import BarChart, Reference
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

from funding_calculator import FundingCalculator
from shock_analyzer import HistoricalShockAnalyzer
from shock_monte_carlo import ShockMonteCarloSimulator

# ページ設定
st.set_page_config(
    page_title="資金繰りリスク分析",
    page_icon="💰",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# カスタムCSS（スマホ対応＋視認性向上）
st.markdown("""
<style>
    /* セクションヘッダーを大きく、色付き背景 */
    .section-header {
        font-size: 1.8rem;
        font-weight: bold;
        padding: 15px;
        margin: 20px 0 15px 0;
        border-radius: 8px;
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        text-align: center;
    }
    
    .section-header-success {
        background: linear-gradient(135deg, #56ab2f 0%, #a8e063 100%);
    }
    
    .section-header-warning {
        background: linear-gradient(135deg, #f093fb 0%, #f5576c 100%);
    }
    
    .section-header-info {
        background: linear-gradient(135deg, #4facfe 0%, #00f2fe 100%);
    }
    
    /* ボタンを大きく */
    .stButton > button {
        width: 100%;
        height: 60px;
        font-size: 1.2rem;
        font-weight: bold;
        border-radius: 10px;
    }
    
    /* ダウンロードボタンを大きく */
    .stDownloadButton > button {
        width: 100%;
        height: 55px;
        font-size: 1.1rem;
        font-weight: bold;
    }
    
    /* 入力フィールドのラベルを大きく */
    .stNumberInput label, .stSelectbox label {
        font-size: 1.1rem;
        font-weight: 600;
    }
    
    /* スマホで列を自動調整 */
    @media (max-width: 768px) {
        .row-widget.stHorizontal {
            flex-direction: column;
        }
    }
</style>
""", unsafe_allow_html=True)

# =============================================================================
# Word出力関数
# =============================================================================
def generate_word_report(params, funding_needs, shock_results, monthly_cf_surplus):
    """Word形式の銀行交渉用レポートを生成"""
    
    if not DOCX_AVAILABLE:
        return None
    
    doc = Document()
    
    # タイトル
    title = doc.add_heading('資金繰り安定化のための長期運転資金申込書', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 作成日
    doc.add_paragraph(f'作成日: {datetime.now().strftime("%Y年%m月%d日")}')
    doc.add_paragraph('')
    
    # 1. 現状の課題
    doc.add_heading('1. 現状の課題', 1)
    
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Light Grid Accent 1'
    
    cells = table.rows[0].cells
    cells[0].text = '項目'
    cells[1].text = '内容'
    
    cells = table.rows[1].cells
    cells[0].text = '現金残高'
    cells[1].text = f'{params["cash_balance"]:,.0f}万円'
    
    cells = table.rows[2].cells
    cells[0].text = '評価'
    cells[1].text = '資金繰り安定化が必要'
    
    doc.add_paragraph('')
    
    # 2. 申込内容（安全水準）
    doc.add_heading('2. 申込内容', 1)
    
    recommended_funding = funding_needs['安全']['funding_amount']
    recommended_repayment = funding_needs['安全']['monthly_repayment_5y']
    
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Grid Accent 1'
    
    cells = table.rows[0].cells
    cells[0].text = '項目'
    cells[1].text = '金額'
    
    cells = table.rows[1].cells
    cells[0].text = '長期運転資金'
    cells[1].text = f'{recommended_funding:,.0f}万円'
    
    cells = table.rows[2].cells
    cells[0].text = '返済期間'
    cells[1].text = '5年（60ヶ月）'
    
    cells = table.rows[3].cells
    cells[0].text = '月次返済額'
    cells[1].text = f'{recommended_repayment:,.1f}万円'
    
    doc.add_paragraph('')
    
    # 3. 導入効果
    doc.add_heading('3. 導入効果', 1)
    doc.add_paragraph(f'• 資金ショート確率: {funding_needs["安全"]["target_probability"]}%以下（安全水準）を達成')
    doc.add_paragraph('• 経営の安定化を実現')
    doc.add_paragraph('')
    
    # 4. 返済計画
    doc.add_heading('4. 返済計画', 1)
    
    gross_profit = params['monthly_sales'] * (1 - params['cost_rate']/100)
    
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Light Grid Accent 1'
    
    rows_data = [
        ('項目', '金額'),
        ('月次粗利平均', f'{gross_profit:,.0f}万円'),
        ('固定費', f'{params["monthly_fixed_cost"]:,.0f}万円'),
        ('CF余剰', f'{monthly_cf_surplus:,.0f}万円'),
        ('月次返済額', f'{recommended_repayment:,.1f}万円'),
        ('返済後のCF余裕', f'{monthly_cf_surplus - recommended_repayment:,.1f}万円')
    ]
    
    for i, (item, value) in enumerate(rows_data):
        cells = table.rows[i].cells
        cells[0].text = item
        cells[1].text = value
    
    doc.add_paragraph('')
    
    # 5. ストレステスト結果
    doc.add_heading('5. ストレステスト結果', 1)
    doc.add_paragraph('歴史的危機シナリオでの耐性:')
    doc.add_paragraph('')
    
    for scenario, result in list(shock_results.items())[:5]:  # 上位5つ
        prob = result['shortage_probability']
        doc.add_paragraph(f'• {scenario}: 資金ショート確率 {prob}%')
    
    doc.add_paragraph('')
    doc.add_paragraph('※リーマンショック級の危機でもリスクを定量化し、対策を立てることが可能です。')
    
    # Wordファイルをバイナリデータとして保存
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer

# =============================================================================
# Excel出力関数
# =============================================================================
def generate_excel_report(params, funding_needs, shock_results, monthly_cf_surplus):
    """Excel形式の詳細レポートを生成"""
    
    if not EXCEL_AVAILABLE:
        return None
    
    wb = openpyxl.Workbook()
    
    # シート1: 必要資金額
    ws1 = wb.active
    ws1.title = "必要資金額"
    
    # ヘッダー
    ws1['A1'] = '目標水準別 必要資金額'
    ws1['A1'].font = Font(size=16, bold=True)
    
    # 列ヘッダー
    headers = ['目標水準', '資金ショート確率', '必要現金残高', '必要資金額', '月次返済額（5年）']
    for col, header in enumerate(headers, 1):
        cell = ws1.cell(row=3, column=col)
        cell.value = header
        cell.font = Font(bold=True)
        cell.fill = PatternFill(start_color='366092', end_color='366092', fill_type='solid')
        cell.font = Font(color='FFFFFF', bold=True)
    
    # データ
    row = 4
    for level in ['安全', '標準', '最低限']:
        result = funding_needs[level]
        ws1.cell(row, 1).value = level
        ws1.cell(row, 2).value = f'{result["target_probability"]}%以下'
        ws1.cell(row, 3).value = result['required_cash']
        ws1.cell(row, 3).number_format = '#,##0'
        ws1.cell(row, 4).value = result['funding_amount']
        ws1.cell(row, 4).number_format = '#,##0'
        ws1.cell(row, 5).value = result['monthly_repayment_5y']
        ws1.cell(row, 5).number_format = '#,##0.0'
        row += 1
    
    # 列幅調整
    ws1.column_dimensions['A'].width = 15
    ws1.column_dimensions['B'].width = 18
    ws1.column_dimensions['C'].width = 18
    ws1.column_dimensions['D'].width = 18
    ws1.column_dimensions['E'].width = 20
    
    # シート2: ストレステスト
    ws2 = wb.create_sheet("ストレステスト")
    
    ws2['A1'] = 'ストレステスト結果'
    ws2['A1'].font = Font(size=16, bold=True)
    
    # 列ヘッダー
    headers = ['シナリオ', '資金ショート確率(%)', '最低残高平均(万円)']
    for col, header in enumerate(headers, 1):
        cell = ws2.cell(row=3, column=col)
        cell.value = header
        cell.font = Font(bold=True)
        cell.fill = PatternFill(start_color='366092', end_color='366092', fill_type='solid')
        cell.font = Font(color='FFFFFF', bold=True)
    
    # データ
    row = 4
    for scenario, result in shock_results.items():
        ws2.cell(row, 1).value = scenario
        ws2.cell(row, 2).value = result['shortage_probability']
        ws2.cell(row, 3).value = result['min_cash_mean']
        ws2.cell(row, 3).number_format = '#,##0.0'
        row += 1
    
    # 列幅調整
    ws2.column_dimensions['A'].width = 20
    ws2.column_dimensions['B'].width = 22
    ws2.column_dimensions['C'].width = 22
    
    # シート3: 基本情報
    ws3 = wb.create_sheet("基本情報")
    
    ws3['A1'] = '入力パラメータ'
    ws3['A1'].font = Font(size=16, bold=True)
    
    row = 3
    param_labels = {
        'monthly_sales': '月次売上高',
        'cash_balance': '現金残高',
        'cost_rate': '売上原価率',
        'monthly_fixed_cost': '月次固定費',
        'sales_volatility': '売上変動率',
        'ar_days': '売掛サイト',
        'ap_days': '支払サイト',
        'inventory_days': '在庫回転期間'
    }
    
    for key, label in param_labels.items():
        ws3.cell(row, 1).value = label
        ws3.cell(row, 1).font = Font(bold=True)
        
        value = params[key]
        if key in ['cost_rate', 'sales_volatility']:
            ws3.cell(row, 2).value = f'{value}%'
        elif key in ['ar_days', 'ap_days', 'inventory_days']:
            ws3.cell(row, 2).value = f'{value}日'
        else:
            ws3.cell(row, 2).value = f'{value:,.0f}万円'
        
        row += 1
    
    ws3.column_dimensions['A'].width = 20
    ws3.column_dimensions['B'].width = 18
    
    # Excelファイルをバイナリデータとして保存
    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    
    return buffer

# =============================================================================
# メインアプリ
# =============================================================================

# タイトル
st.markdown("# 💰 資金繰りリスク可視化ツール")
st.markdown("### 中小企業経営者向け - 5分で資金繰りリスクを診断")

st.info("""
📋 **このツールの使い方**

1️⃣ 財務数値を入力（8項目のみ）  
2️⃣ シミュレーション実行ボタンをクリック  
3️⃣ 長期運転資金 → ストレステスト → レポート出力
""")

# 実行ボタン
st.markdown("---")
run_simulation = st.button(
    "🚀 シミュレーション実行（10,000回）",
    type="primary",
    use_container_width=True
)
st.markdown("---")

# 基本情報入力
st.markdown('<div class="section-header">📋 ステップ1: 基本情報の入力</div>', 
            unsafe_allow_html=True)

col1, col2 = st.columns(2)

with col1:
    industry = st.selectbox(
        "業種",
        options=['建設業', '製造業', '卸・小売業', '運輸・物流業', 'サービス・その他']
    )

with col2:
    employee_count = st.selectbox(
        "従業員数",
        options=['30人以下', '30人超～50人以下', '50人超～100人以下', 
                '100人超～300人以下', '300人超']
    )

service_sub = None
if industry == 'サービス・その他':
    service_sub = st.selectbox(
        "業態を選択",
        options=['IT・システム開発', '医療・介護', '広告・マーケティング', 'その他']
    )

# 財務数値入力
st.markdown('<div class="section-header">💵 ステップ2: 財務数値の入力</div>', 
            unsafe_allow_html=True)

col1, col2 = st.columns(2)
with col1:
    monthly_sales = st.number_input("月次売上高（万円）", min_value=0, value=1000, step=100)
with col2:
    cash_balance = st.number_input("現金残高（万円）", min_value=0, value=800, step=100)

col1, col2 = st.columns(2)
with col1:
    cost_rate = st.number_input("売上原価率（%）", min_value=0.0, max_value=100.0, value=65.0, step=1.0)
with col2:
    monthly_fixed_cost = st.number_input("月次固定費（万円）", min_value=0, value=320, step=10)

col1, col2 = st.columns(2)
with col1:
    sales_volatility = st.number_input("売上変動率（%）", min_value=0.0, max_value=50.0, value=15.0, step=1.0)
with col2:
    ar_days = st.number_input("売掛サイト（日）", min_value=0, max_value=180, value=45, step=5)

col1, col2 = st.columns(2)
with col1:
    ap_days = st.number_input("支払サイト（日）", min_value=0, max_value=180, value=35, step=5)
with col2:
    inventory_days = st.number_input("在庫回転期間（日）", min_value=0, max_value=180, value=30, step=5)

# シミュレーション実行
if run_simulation:
    params = {
        'monthly_sales': monthly_sales,
        'sales_volatility': sales_volatility,
        'cash_balance': cash_balance,
        'cost_rate': cost_rate,
        'monthly_fixed_cost': monthly_fixed_cost,
        'ar_days': ar_days,
        'ap_days': ap_days,
        'inventory_days': inventory_days,
        'existing_debt_repayment': 0
    }
    
    # Phase 3: 長期運転資金
    st.markdown('<div class="section-header section-header-success">💰 Phase 3: 長期運転資金の必要額</div>', 
                unsafe_allow_html=True)
    
    with st.spinner('必要資金額を算出中...'):
        calculator = FundingCalculator(n_simulations=10000)
        funding_needs = calculator.calculate_funding_needs(params)
        monthly_cf_surplus = calculator.calculate_monthly_cf_surplus(params)
    
    st.subheader("📊 目標水準別の必要資金額")
    
    for level in ['安全', '標準', '最低限']:
        result = funding_needs[level]
        
        if result['funding_amount'] == 0:
            color, icon = "#d4edda", "✅"
        elif result['funding_amount'] < monthly_sales:
            color, icon = "#fff3cd", "⚠️"
        else:
            color, icon = "#f8d7da", "🔴"
        
        st.markdown(f"""
        <div style="background-color: {color}; padding: 20px; border-radius: 10px; margin: 10px 0;">
            <h3>{icon} {level}水準（資金ショート確率 {result['target_probability']}%以下）</h3>
            <p style="font-size: 1.3rem; font-weight: bold;">
                必要資金額: {result['funding_amount']:,.0f}万円
            </p>
            <p style="font-size: 1.1rem;">
                必要現金残高: {result['required_cash']:,.0f}万円<br>
                月次返済額（5年）: {result['monthly_repayment_5y']:,.1f}万円
            </p>
        </div>
        """, unsafe_allow_html=True)
    
    # Phase 2: ストレステスト
    st.markdown('<div class="section-header section-header-warning">🔥 Phase 2: 歴史的ストレステスト</div>', 
                unsafe_allow_html=True)
    
    with st.spinner('ストレステスト実行中...'):
        simulator = ShockMonteCarloSimulator(n_simulations=10000)
        shock_results = simulator.compare_all_shocks(params, industry, service_sub)
    
    st.subheader("📊 ショック別資金ショート確率")
    
    for scenario, result in shock_results.items():
        prob = result['shortage_probability']
        
        if prob >= 50:
            color, icon = "#f8d7da", "🔴"
        elif prob >= 30:
            color, icon = "#fff3cd", "⚠️"
        else:
            color, icon = "#d4edda", "✅"
        
        st.markdown(f"""
        <div style="background-color: {color}; padding: 15px; border-radius: 8px; margin: 8px 0;">
            <h4>{icon} {scenario}</h4>
            <p style="font-size: 1.2rem; font-weight: bold;">
                資金ショート確率: {prob}%
            </p>
        </div>
        """, unsafe_allow_html=True)
    
    # レポート出力
    st.markdown('<div class="section-header section-header-info">📄 レポート出力</div>', 
                unsafe_allow_html=True)
    
    st.info("**銀行交渉用レポートをダウンロードできます**")
    
    col1, col2 = st.columns(2)
    
    # Word出力
    with col1:
        if DOCX_AVAILABLE:
            word_buffer = generate_word_report(params, funding_needs, shock_results, monthly_cf_surplus)
            if word_buffer:
                st.download_button(
                    label="📝 Wordレポート",
                    data=word_buffer,
                    file_name=f"資金繰り分析_{datetime.now().strftime('%Y%m%d')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
        else:
            st.error("python-docxがインストールされていません")
    
    # Excel出力
    with col2:
        if EXCEL_AVAILABLE:
            excel_buffer = generate_excel_report(params, funding_needs, shock_results, monthly_cf_surplus)
            if excel_buffer:
                st.download_button(
                    label="📊 Excelレポート",
                    data=excel_buffer,
                    file_name=f"資金繰り分析_{datetime.now().strftime('%Y%m%d')}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True
                )
        else:
            st.error("openpyxlがインストールされていません")

st.markdown("---")
st.markdown("""
<div style="text-align: center; color: #666; font-size: 0.9rem;">
💡 操作方法がわからない場合は、各項目の「?」マークをクリックしてください<br>
📱 スマートフォンでも快適にご利用いただけます
</div>
""", unsafe_allow_html=True)
